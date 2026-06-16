import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()


title = doc.add_heading("Rapport de Révision : Projet de Fin d'Études (NetWatch)", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Objet : Audit technique et typographique du rapport de PFE."
).bold = True
doc.add_paragraph(
    "Le rapport est globalement d'excellente qualité, structuré de manière cohérente, et met en évidence un travail technique solide. Cependant, quelques corrections sont nécessaires avant la soumission finale pour garantir un niveau de professionnalisme irréprochable."
)

p_alert = doc.add_paragraph("Note générale : ")
p_alert.add_run(
    "Les erreurs soulevées ici sont principalement des conflits de terminologie technique, des fautes de frappe et des problèmes d'espacement souvent causés par le moteur de rendu (LaTeX)."
)
p_alert.style = "Intense Quote"


doc.add_heading("1. Conflit Technique Majeur 🚨", level=1)
p1 = doc.add_paragraph(
    "Stack MERN vs NestJS (Chapitre 3, Section 3.1) : ", style="List Bullet"
)
p1.runs[0].bold = True
p1.add_run(
    "Il est mentionné l'utilisation de la stack MERN (spécifiquement React + NestJS + MongoDB + Node.js). C'est un conflit technique. L'acronyme \"MERN\" désigne précisément MongoDB, "
)
p1.add_run("Express.js").bold = True
p1.add_run(", React et Node.js. Puisque l'API utilise ")
p1.add_run("NestJS").bold = True
p1.add_run(" (qui remplace Express), ce n'est plus une stack MERN.\n")
p1.add_run("➔ Correction suggérée : ").bold = True
r = p1.add_run(
    '"une stack JavaScript/TypeScript moderne (React, NestJS, MongoDB et Node.js)"'
)
r.font.color.rgb = RGBColor(0, 128, 0)
p1.add_run(' ou retirer simplement le terme "MERN".')


doc.add_heading("2. Fautes d'Orthographe et de Frappe 📝", level=1)
p2_1 = doc.add_paragraph(
    "Section 3.4.2.3 (Journalisation des événements) : ", style="List Bullet"
)
p2_1.runs[0].bold = True
p2_1.add_run(
    'Il est écrit : "Cet écran assure une raçabilité des actions". Il manque un "t".\n'
)
p2_1.add_run("➔ Correction : ").bold = True
r = p2_1.add_run('"traçabilité".')
r.font.color.rgb = RGBColor(0, 128, 0)

p2_2 = doc.add_paragraph(
    "Tableau 2.9 (Description des classes, N° 10) : ", style="List Bullet"
)
p2_2.runs[0].bold = True
p2_2.add_run("La méthode est nommée \"envoi_maile()\", avec un 'e' en trop à la fin.\n")
p2_2.add_run("➔ Correction : ").bold = True
r = p2_2.add_run('"envoi_mail()" ou "envoyer_mail()".')
r.font.color.rgb = RGBColor(0, 128, 0)

p2_3 = doc.add_paragraph(
    "Tableau 2.9 (Description des classes, N° 8) : ", style="List Bullet"
)
p2_3.runs[0].bold = True
p2_3.add_run('Le nom de la classe est "PortailCapt", ce qui semble tronqué.\n')
p2_3.add_run("➔ Correction suggérée : ").bold = True
r = p2_3.add_run('La renommer en "PortailCaptif".')
r.font.color.rgb = RGBColor(0, 128, 0)

p2_4 = doc.add_paragraph(
    "Tableau 2.9 (Description des classes, N° 6) : ", style="List Bullet"
)
p2_4.runs[0].bold = True
p2_4.add_run(
    'Il est écrit : "orchestre par composition () les modules". Il y a des parenthèses vides.\n'
)
p2_4.add_run("➔ Correction : ").bold = True
r = p2_4.add_run(
    "Soit ajouter un symbole ou une précision à l'intérieur, soit supprimer ces parenthèses."
)
r.font.color.rgb = RGBColor(0, 128, 0)


doc.add_heading("3. Incohérences de Traduction (Français / Anglais) 🌍", level=1)
p3 = doc.add_paragraph(
    "Section 3.2.1.1 (Description des Couches) : ", style="List Bullet"
)
p3.runs[0].bold = True
p3.add_run(
    'Bien que le rapport soit rédigé en français, les noms des couches architecturales sont restés en anglais : "Presentation Layer", "Business Logic Layer", "Data Layer", "Auth Validation", "External Comm".\n'
)
p3.add_run("➔ Correction suggérée : ").bold = True
p3.add_run("Les traduire en français, par exemple : ")
r = p3.add_run(
    '"Couche de présentation (Presentation Layer)", "Couche logique métier", "Couche de données".'
)
r.font.color.rgb = RGBColor(0, 128, 0)


doc.add_heading(
    "4. Erreurs de Typographie et d'Espacements (Liées à LaTeX) 🔍", level=1
)
doc.add_paragraph(
    "De nombreux espaces entre les mots ont disparu. C'est un problème très courant lors de la compilation LaTeX (souvent dû à des macros comme \\textbf{} ou \\url{} qui avalent l'espace suivant)."
)

p4_1 = doc.add_paragraph('Page "Autorisation de dépôt" :', style="List Bullet")
p4_1.runs[0].bold = True
sub_p4_1_1 = doc.add_paragraph(
    "M.Tarek Dhokkar ➔ M. Tarek Dhokkar", style="List Bullet 2"
)
sub_p4_1_2 = doc.add_paragraph(
    "l’étudianteSelmi Salsabil ➔ l'étudiante Selmi Salsabil", style="List Bullet 2"
)
sub_p4_1_3 = doc.add_paragraph(
    "MmeHenda Boudegga ➔ Mme Henda Boudegga", style="List Bullet 2"
)

p4_2 = doc.add_paragraph("Dans le corps du texte :", style="List Bullet")
p4_2.runs[0].bold = True
sub_p4_2_1 = doc.add_paragraph("sécurisésJWT ➔ sécurisés JWT", style="List Bullet 2")
sub_p4_2_2 = doc.add_paragraph("Mode Sombreet ➔ Mode Sombre et", style="List Bullet 2")
sub_p4_2_3 = doc.add_paragraph("APIResend ➔ API Resend", style="List Bullet 2")
sub_p4_2_4 = doc.add_paragraph("Socket.IOet ➔ Socket.IO et", style="List Bullet 2")
sub_p4_2_5 = doc.add_paragraph(
    "JavaScript:Langagedebase... (Tableau 3.1) ➔ Remettre les espaces.",
    style="List Bullet 2",
)

p4_3 = doc.add_paragraph("Bibliographie : ", style="List Bullet")
p4_3.runs[0].bold = True
p4_3.add_run(
    'Les URLs sont souvent collées au mot précédent (ex: "marketing.https://..."). Il faut ajouter un espace avant chaque URL.'
)


doc.add_heading('5. Incohérence "Table" vs "Tableau" 📊', level=1)
p5 = doc.add_paragraph(
    'Dans le texte, il est écrit "Le tableau 1.1 présente...", mais au-dessus des tableaux, la légende générée affiche "Table 1.1 : ...".\n',
    style="List Bullet",
)
p5.add_run("➔ Correction : ").bold = True
p5.add_run(
    'En français, le terme standard est "Tableau". Il faut configurer le package babel avec l\'option french dans LaTeX pour corriger ces légendes générées automatiquement.'
)


doc.add_heading("6. Légendes des Figures 🖼️", level=1)
doc.add_paragraph(
    'Figure 1.1 : "logo sotupub" ➔ À corriger en "Logo SOTUPUB".', style="List Bullet"
)
doc.add_paragraph(
    'Figure 2.1 : "logo uml" ➔ À corriger en "Logo UML".', style="List Bullet"
)
doc.add_paragraph(
    'Figure 2.10 : "Diagramme de classes global ." ➔ Il y a un espace en trop avant le point final. De plus, il est d\'usage de ne pas mettre de point à la fin des légendes de figures.',
    style="List Bullet",
)


doc.save("Rapport_Correction.docx")
print("Document DOCX créé avec succès.")
